#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1, spacing_before=12, spacing_after=6):
    """Add a styled heading"""
    heading = doc.add_paragraph()
    heading_run = heading.add_run(text)
    
    if level == 1:
        heading_run.font.size = Pt(18)
        heading_run.font.bold = True
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
    else:
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
    
    heading_run.font.name = 'Arial'
    heading.paragraph_format.line_spacing = 1.5
    heading.paragraph_format.space_before = Pt(spacing_before)
    heading.paragraph_format.space_after = Pt(spacing_after)
    
    return heading

def add_paragraph_text(doc, text, font_size=12, spacing_line=1.5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Add a paragraph with consistent formatting"""
    paragraph = doc.add_paragraph(text)
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(font_size)
    paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = spacing_line
    return paragraph

# Create document
doc = Document()

# ============ PORTADA ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("MANUAL DE USUARIO")
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.name = 'Arial'
title.paragraph_format.space_after = Pt(6)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("CuidaLink")
subtitle_run.font.size = Pt(24)
subtitle_run.font.name = 'Arial'
subtitle_run.font.bold = True
subtitle.paragraph_format.space_after = Pt(12)

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc_run = desc.add_run("Plataforma Digital de Asistencia\npara Cuidadores de Personas con Alzheimer")
desc_run.font.size = Pt(14)
desc_run.font.name = 'Arial'
desc_run.font.italic = True
desc.paragraph_format.space_after = Pt(24)

# Add space
for _ in range(6):
    doc.add_paragraph()

# Información
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info.add_run("Versión 1.0\nMarzo de 2026")
info_run.font.size = Pt(12)
info_run.font.name = 'Arial'

doc.add_page_break()

# ============ ÍNDICE ============
add_heading(doc, "ÍNDICE", 1, spacing_before=0, spacing_after=12)

toc_items = [
    "1. Introducción",
    "2. Requisitos del Sistema",
    "3. Instalación y Configuración Inicial",
    "4. Guía de Uso - Rol Cuidadora",
    "   4.1. Login y Acceso",
    "   4.2. Dashboard de Tareas",
    "   4.3. Verificación de Llegada (Check-in)",
    "   4.4. Medicación (Pastilla)",
    "   4.5. Comida",
    "   4.6. Paseo (Tracking GPS)",
    "   4.7. Siesta",
    "   4.8. Chat con Familia",
    "   4.9. Notas del Familiar",
    "   4.10. Valoración Diaria",
    "   4.11. Calendario",
    "5. Guía de Uso - Rol Familiar",
    "   5.1. Login y Acceso",
    "   5.2. Panel de Supervisión",
    "   5.3. Gestión de Medicación",
    "   5.4. Envío de Notas",
    "   5.5. Chat",
    "   5.6. Informes Semanales",
    "6. Funcionalidades Comunes",
    "   6.1. Mapa y Zona Segura",
    "   6.2. Localización en Tiempo Real",
    "   6.3. Detección de Caídas",
    "7. Preguntas Frecuentes (FAQ)",
    "8. Troubleshooting",
    "9. Contacto y Soporte",
]

for item in toc_items:
    toc_para = doc.add_paragraph(item)
    toc_para.paragraph_format.left_indent = Inches(0.25) if item.startswith("   ") else Inches(0)
    for run in toc_para.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    toc_para.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ============ 1. INTRODUCCIÓN ============
add_heading(doc, "1. INTRODUCCIÓN", 1, spacing_before=0, spacing_after=12)

add_paragraph_text(doc,
"Bienvenido a CuidaLink, una aplicación móvil diseñada para mejorar la comunicación y transparencia en el cuidado de personas con Alzheimer. Este manual le guiará paso a paso sobre cómo utilizar todas las funcionalidades de la aplicación, tanto si es cuidador/a profesional como si es familiar del paciente.")

add_paragraph_text(doc,
"CuidaLink conecta a cuidadores profesionales con familias, proporcionando:")

features_intro = [
    "Verificación fotográfica de medicación y comidas",
    "Geolocalización en tiempo real durante paseos",
    "Comunicación bidireccional (chat, texto, imagen, audio)",
    "Registro automático de actividades diarias",
    "Sistema de valoración de calidad del cuidado",
    "Detección automática de caídas",
    "Acceso offline-first (funciona sin internet)",
    "Reportes y análisis semanales",
]

for feature in features_intro:
    f = doc.add_paragraph(feature, style='List Bullet')
    for run in f.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    f.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ============ 2. REQUISITOS DEL SISTEMA ============
add_heading(doc, "2. REQUISITOS DEL SISTEMA", 1, spacing_before=0, spacing_after=12)

add_heading(doc, "2.1. Dispositivos Compatible", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"CuidaLink está disponible para iOS y Android. A continuación se detallan los requisitos mínimos:")

reqs = [
    ("iOS", "Versión 13 o superior. iPhone 6s o más reciente."),
    ("Android", "Versión 8 (API 26) o superior. Memoria mínima: 2 GB RAM."),
]

for sistema, requisito in reqs:
    p = doc.add_paragraph(f"{sistema}: {requisito}", style='List Bullet')
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5

add_heading(doc, "2.2. Conexión y Sensores", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Para funcionalidad completa, se recomienda:")

specs = [
    "Conexión Wi-Fi o datos móviles (4G/5G)",
    "GPS activado para geolocalización",
    "Cámara funcional para verificación fotográfica",
    "Acelerómetro (para detección de caídas)",
    "Micrófono (para notas de audio)",
    "Almacenamiento disponible: mínimo 500 MB",
]

for spec in specs:
    s = doc.add_paragraph(spec, style='List Bullet')
    for run in s.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    s.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ============ 3. INSTALACIÓN ============
add_heading(doc, "3. INSTALACIÓN Y CONFIGURACIÓN INICIAL", 1, spacing_before=0, spacing_after=12)

add_heading(doc, "3.1. Descarga e Instalación", 2, spacing_before=6, spacing_after=6)

add_paragraph_text(doc,
"PARA ANDROID:", font_size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_paragraph_text(doc,
"1. Abre Google Play Store en tu dispositivo Android.\n2. Busca 'CuidaLink' en la barra de búsqueda.\n3. Haz clic en 'Instalar' y espera a que se descargue.\n4. Una vez instalado, haz clic en 'Abrir'.")

add_paragraph_text(doc,
"PARA iOS:", font_size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_paragraph_text(doc,
"1. Abre App Store en tu dispositivo iOS.\n2. Busca 'CuidaLink' en la pestaña Buscar.\n3. Haz clic en 'Obtener' y completa la autenticación con Face ID o Touch ID.\n4. Una vez instalada, haz clic en 'Abrir'.")

add_heading(doc, "3.2. Primer Acceso y Registro", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"1. Al abrir CuidaLink por primera vez, verás la pantalla de bienvenida.\n2. Selecciona tu rol: 'Cuidadora' o 'Familiar'.\n3. Introduce tu teléfono y PIN (código de acceso de 4 dígitos).\n4. Haz clic en 'Entrar'.\n5. Si es tu primer acceso, se te pedirá confirmar permisos de cámara, ubicación y notificaciones. Haz clic en 'Permitir' para cada uno.\n6. ¡Bienvenido! Ya tienes acceso a CuidaLink.")

add_paragraph_text(doc,
"NOTA: El PIN se proporciona por separado por el administrador del sistema.",
font_size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)

doc.add_page_break()

# ============ 4. GUÍA CUIDADORA ============
add_heading(doc, "4. GUÍA DE USO - ROL CUIDADORA", 1, spacing_before=0, spacing_after=12)

add_heading(doc, "4.1. Login y Acceso", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Pasos para acceder a tu cuenta como cuidadora:")
steps = [
    "Abre la aplicación CuidaLink.",
    "En la pantalla de selección de rol, haz clic en 'Cuidadora'.",
    "Introduce tu número de teléfono (formato: 6XXXXXXXX).",
    "Introduce tu PIN de 4 dígitos.",
    "Haz clic en 'Entrar'.",
    "Serás redirigido al Dashboard de Tareas.",
]
for i, step in enumerate(steps, 1):
    s = doc.add_paragraph(f"{i}. {step}")
    for run in s.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    s.paragraph_format.line_spacing = 1.5

add_heading(doc, "4.2. Dashboard de Tareas", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"El Dashboard es tu pantalla principal donde puedes ver todas las tareas del día. Muestra:")
dashboard_items = [
    "Barra de progreso: Porcentaje de tareas completadas (0-100%).",
    "Lista de 5 tareas diarias: Llegada (🚪), Medicación (💊), Paseo (🚶), Comida (🍽️), Siesta (😴).",
    "Estado de cada tarea: Pendiente (gris), En curso (azul), Completada (verde).",
    "Notas del familiar: Mensajes urgentes o recordatorios especiales.",
    "Botón de Confetti: Celebración animada al completar todas las tareas.",
]
for item in dashboard_items:
    ds = doc.add_paragraph(item, style='List Bullet')
    for run in ds.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    ds.paragraph_format.line_spacing = 1.5

add_heading(doc, "4.3. Verificación de Llegada (Check-in)", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Esta es la primera tarea del día. Registra tu llegada con una selfie:")
checkin_steps = [
    "Haz clic en la tarjeta 'Llegada' en el Dashboard.",
    "La cámara frontral se abrirá automáticamente.",
    "Tómate una selfie clara (asegúrate de que la iluminación sea buena).",
    "La aplicación verificará automáticamente tu ubicación con GPS.",
    "Haz clic en 'Confirmar' para registrar tu llegada.",
    "Verás un mensaje de éxito y la tarea se marcará como completada.",
]
for i, step in enumerate(checkin_steps, 1):
    cs = doc.add_paragraph(f"{i}. {step}")
    for run in cs.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    cs.paragraph_format.line_spacing = 1.5

add_heading(doc, "4.4. Medicación (Pastilla)", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Verifica la administración de medicación fotográficamente:")
med_steps = [
    "Haz clic en la tarjeta 'Pastilla' en el Dashboard.",
    "Se abrirá la cámara trasera.",
    "Fotografía el medicamento junto con la etiqueta visible.",
    "La aplicación realizará OCR simulado (lectura de la etiqueta).",
    "Revisa la información extraída: nombre, dosis, hora.",
    "Haz clic en 'Confirmar administración'.",
]
for i, step in enumerate(med_steps, 1):
    ms = doc.add_paragraph(f"{i}. {step}")
    for run in ms.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    ms.paragraph_format.line_spacing = 1.5

add_heading(doc, "4.5. Comida", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Documenta la comida con una fotografía:")
food_steps = [
    "Haz clic en la tarjeta 'Comida' en el Dashboard.",
    "Se abrirá la cámara trasera.",
    "Fotografía el plato de comida.",
    "Selecciona el nivel de apetito del paciente: Bueno (😊), Normal (😐), Bajo (😞).",
    "Añade un comentario opcional sobre la comida (sabor, cantidad, etc.).",
    "Haz clic en 'Guardar Comida'.",
]
for i, step in enumerate(food_steps, 1):
    fs = doc.add_paragraph(f"{i}. {step}")
    for run in fs.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    fs.paragraph_format.line_spacing = 1.5

add_heading(doc, "4.6. Paseo (Tracking GPS)", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Realiza el seguimiento de los paseos en tiempo real:")
walk_steps = [
    "Haz clic en la tarjeta 'Paseo' en el Dashboard.",
    "Se abrirá la pantalla de Tracking con un mapa.",
    "La aplicación comienza a registrar tu ubicación cada 10 segundos.",
    "Durante el paseo, verás en tiempo real: distancia recorrida, duración, velocidad.",
    "La zona segura (círculo) aparecerá en el mapa. Si abandonas esta zona, se enviará una alerta a la familia.",
    "Haz clic en 'Pausar' si necesitas detener el tracking temporalmente.",
    "Haz clic en 'Finalizar Paseo' al terminar.",
    "Verás un resumen: distancia total, duración, ruta.",
]
for i, step in enumerate(walk_steps, 1):
    ws = doc.add_paragraph(f"{i}. {step}")
    for run in ws.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    ws.paragraph_format.line_spacing = 1.5

add_heading(doc, "4.7. Siesta", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Registra el tiempo de descanso del paciente:")
siesta_steps = [
    "Haz clic en la tarjeta 'Siesta' en el Dashboard.",
    "Se abrirá un temporizador.",
    "Selecciona la duración estimada de la siesta (ej: 30 minutos, 1 hora).",
    "Haz clic en 'Iniciar Siesta'.",
    "El temporizador cuenta hacia atrás. Durante este tiempo, no puedes cerrar la pantalla.",
    "Cuando suene la alarma, verás una notificación.",
    "Haz clic en 'Siesta completada' para registrarla.",
]
for i, step in enumerate(siesta_steps, 1):
    ss = doc.add_paragraph(f"{i}. {step}")
    for run in ss.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    ss.paragraph_format.line_spacing = 1.5

add_heading(doc, "4.8. Chat con Familia", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Comunícate directamente con la familia del paciente:")
chat_steps = [
    "Toca el icono de 'Chat' en la barra inferior.",
    "Verás el historial de conversaciones.",
    "Haz clic en la conversación con la familia.",
    "Escribe tu mensaje en la caja de texto inferior.",
    "Opciones de envío: Texto normal, Emoji, Imagen (từ cámara o galería), Audio (graba un mensaje).",
    "Haz clic en 'Enviar' (ícono de flecha).",
]
for i, step in enumerate(chat_steps, 1):
    chs = doc.add_paragraph(f"{i}. {step}")
    for run in chs.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    chs.paragraph_format.line_spacing = 1.5

add_heading(doc, "4.9. Notas del Familiar", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Lee las notas y recordatorios del familiar:")
notes_steps = [
    "Toca el icono de 'Notas' en la barra inferior.",
    "Verás todas las notas ordenadas por fecha (más recientes primero).",
    "Las notas urgentes aparecen con un ícono 🚨.",
    "Haz clic en una nota para leerla completa.",
    "Se registra automáticamente el timestamp de lectura.",
    "Haz clic en '✓ Leído' para confirmar la lectura.",
]
for i, step in enumerate(notes_steps, 1):
    ns = doc.add_paragraph(f"{i}. {step}")
    for run in ns.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    ns.paragraph_format.line_spacing = 1.5

add_heading(doc, "4.10. Valoración Diaria", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Al finalizar tu jornada, valora la calidad del cuidado prestado:")
rating_steps = [
    "Haz clic en la tarjeta 'Valoración' en el Dashboard.",
    "Verás una escala de 1 a 5 estrellas.",
    "Selecciona tu valoración (1=Malo, 5=Excelente).",
    "Escribe un comentario opcional sobre la jornada.",
    "Haz clic en 'Enviar Valoración'.",
    "Tu puntuación se registra y se envía a la familia.",
]
for i, step in enumerate(rating_steps, 1):
    rs = doc.add_paragraph(f"{i}. {step}")
    for run in rs.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    rs.paragraph_format.line_spacing = 1.5

add_heading(doc, "4.11. Calendario", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Visualiza tu historial de actividades:")
cal_steps = [
    "Toca el icono de 'Calendario' en la barra inferior.",
    "Verás una vista mensual.",
    "Los días con actividades completadas aparecen con un punto verde.",
    "Haz clic en un día para ver el detalle de actividades.",
    "Puedes filtrar por tipo de actividad usando los filtros en la parte superior.",
]
for i, step in enumerate(cal_steps, 1):
    cls = doc.add_paragraph(f"{i}. {step}")
    for run in cls.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    cls.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ============ 5. GUÍA FAMILIAR ============
add_heading(doc, "5. GUÍA DE USO - ROL FAMILIAR", 1, spacing_before=0, spacing_after=12)

add_heading(doc, "5.1. Login y Acceso", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Pasos para acceder como familiar:")
fam_login = [
    "Abre la aplicación CuidaLink.",
    "En la pantalla de selección de rol, haz clic en 'Familiar'.",
    "Introduce tu correo electrónico.",
    "Introduce tu contraseña.",
    "Haz clic en 'Entrar'.",
    "Serás redirigido al Panel de Supervisión.",
]
for i, step in enumerate(fam_login, 1):
    fl = doc.add_paragraph(f"{i}. {step}")
    for run in fl.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    fl.paragraph_format.line_spacing = 1.5

add_heading(doc, "5.2. Panel de Supervisión", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Visualiza el estado actual del cuidado:")
panel_items = [
    "Información del paciente: Nombre, edad, medicamentos activos.",
    "Estado de hoy: Tareas completadas, tareas pendientes, actividad en tiempo real.",
    "Línea temporal: Cronología de eventos del día (llegada, medicación, comida, etc.).",
    "Ubicación actual: Mapa con la ubicación en tiempo real del paciente.",
    "Notificaciones: Alertas de eventos importantes (caída detectada, salida de zona segura, etc.).",
    "Rating: Valoración Media del cuidador (últimos 7 días).",
]
for item in panel_items:
    pi = doc.add_paragraph(item, style='List Bullet')
    for run in pi.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    pi.paragraph_format.line_spacing = 1.5

add_heading(doc, "5.3. Gestión de Medicación", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Administra los medicamentos del paciente:")
med_mgmt = [
    "Toca el icono de 'Medicación' en la barra inferior.",
    "Verás la lista de medicamentos activos.",
    "Haz clic en '+' para añadir un nuevo medicamento.",
    "Completa: nombre, dosis, horarios, frecuencia.",
    "Haz clic en 'Guardar Medicamento'.",
    "Para editar: Haz clic en el medicamento → 'Editar' → Realiza cambios → 'Guardar'.",
    "Para eliminar: Desliza hacia la izquierda → 'Eliminar'.",
    "Verás un gráfico de adherencia semanal (% de dosis tomadas en horario).",
]
for i, step in enumerate(med_mgmt, 1):
    mm = doc.add_paragraph(f"{i}. {step}")
    for run in mm.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    mm.paragraph_format.line_spacing = 1.5

add_heading(doc, "5.4. Envío de Notas", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Envía recordatorios y mensajes al cuidador:")
send_notes = [
    "Toca el icono de 'Notas' en la barra inferior.",
    "Haz clic en '+' para crear una nueva nota.",
    "Escribe el contenido (ej: 'Recordar que hoy hay cita al médico a las 16:30').",
    "Selecciona si es normal o urgente (marca '🚨 Urgente' si es importante).",
    "Haz clic en 'Enviar'.",
    "El cuidador recibirá una notificación push.",
    "Puedes ver el estado de lectura: 'No leído' / 'Leído a las [HH:MM]'.",
]
for i, step in enumerate(send_notes, 1):
    sn = doc.add_paragraph(f"{i}. {step}")
    for run in sn.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    sn.paragraph_format.line_spacing = 1.5

add_heading(doc, "5.5. Chat", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Comunícate en tiempo real con el cuidador:")
chat_fam = [
    "Toca el icono de 'Chat' en la barra inferior.",
    "Abrirá la conversación con el cuidador.",
    "Escribe tu mensaje.",
    "Puedes enviar: Texto, Emoji, Imagen, o Audio.",
    "Los mensajes se sincronizan instantáneamente.",
    "Verás el estado de lectura (✓ Entregado, ✓✓ Leído).",
]
for i, step in enumerate(chat_fam, 1):
    cf = doc.add_paragraph(f"{i}. {step}")
    for run in cf.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    cf.paragraph_format.line_spacing = 1.5

add_heading(doc, "5.6. Informes Semanales", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Analiza el desempeño semanal del cuidador:")
reports = [
    "Toca el icono de 'Informes' en la barra inferior.",
    "Selecciona la semana que quieres analizar (por defecto, la semana actual).",
    "Verás gráficos con:",
    "   - Porcentaje de tareas completadas por día",
    "   - Tiempo total de paseos",
    "   - Adherencia medicamentosa",
    "   - Rating promedio del cuidador",
    "   - Incidents reportados (caídas, alertas de zona segura)",
    "Haz clic en 'Exportar a PDF' para descargar el informe.",
]
for i, item in enumerate(reports, 1):
    if item.startswith("   "):
        rpt = doc.add_paragraph(item, style='List Bullet')
    else:
        rpt = doc.add_paragraph(f"{i}. {item}")
    for run in rpt.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    rpt.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ============ 6. FUNCIONALIDADES COMUNES ============
add_heading(doc, "6. FUNCIONALIDADES COMUNES", 1, spacing_before=0, spacing_after=12)

add_heading(doc, "6.1. Mapa y Zona Segura", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"El mapa muestra la ubicación del paciente y los lugares de interés cercanos (farmacias, hospitales, parques, cafeterías, etc.):")
map_info = [
    "La zona segura aparece como un círculo azul. Se puede ajustar el radio.",
    "Si el paciente abandona esta zona, se envía una alerta a la familia.",
    "Los puntos de interés (POIs) se marcan con emojis: 💊 (farmacia), 🏥 (hospital), 🌳 (parque), ☕ (café).",
    "Haz clic en un POI para ver detalles: nombre, dirección, teléfono, horario.",
    "Puedes habilitar/deshabilitar categorías con los filtros.",
]
for item in map_info:
    mi = doc.add_paragraph(item, style='List Bullet')
    for run in mi.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    mi.paragraph_format.line_spacing = 1.5

add_heading(doc, "6.2. Localización en Tiempo Real", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Ambos roles pueden ver la ubicación actual:")
loc_info = [
    "Cuidadora: Durante un paseo, tu ubicación se registra cada 10 segundos.",
    "Familia: Puede ver la ubicación actual en el panel de supervisión.",
    "Para mayor privacidad, la ubicación se borra después de 7 días.",
    "La precisión es típicamente de 10-20 metros en áreas urbanas.",
]
for item in loc_info:
    li = doc.add_paragraph(item, style='List Bullet')
    for run in li.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    li.paragraph_format.line_spacing = 1.5

add_heading(doc, "6.3. Detección de Caídas", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"La aplicación detecta automáticamente caídas usando el acelerómetro del dispositivo:")
fall_info = [
    "Sistema de doble confirmación: Detecta movimiento de caída + solicita confirmación al usuario.",
    "Si se detecta una caída, la aplicación emite una alarma sonora.",
    "El usuario tiene 30 segundos para confirmar 'Estoy bien' en la pantalla.",
    "Si no confirma, se envía automáticamente una alerta a la familia.",
    "La familia recibirá: ubicación exacta, mapa, opción para llamar al 112.",
    "NOTA: Esta es una capa de seguridad adicional, no reemplaza la atención profesional.",
]
for item in fall_info:
    fi = doc.add_paragraph(item, style='List Bullet')
    for run in fi.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    fi.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ============ 7. FAQ ============
add_heading(doc, "7. PREGUNTAS FRECUENTES (FAQ)", 1, spacing_before=0, spacing_after=12)

faqs = [
    ("¿Qué hago si olvido mi PIN o contraseña?", "Haz clic en 'Olvidé mi contraseña' en la pantalla de login. Recibirás un email con instrucciones para resetear tu acceso. Contacta al soporte si no recibes el email."),
    ("¿Funciona CuidaLink sin internet?", "Sí, la aplicación está diseñada para funcionar offline. Las tareas se registran localmente y se sincronizan automáticamente cuando recuperas conectividad."),
    ("¿Cuánto almacenamiento necesita la app?", "Aproximadamente 500 MB. La mayoría del espacio se usa para el historial de actividades y fotos. Puedes limpiar archivos antiguos manualmente."),
    ("¿Cómo elimino una foto que subí?", "Cuidadora: En la pantalla de revisión de la foto, haz clic en 'Eliminar' antes de confirmar. Familia: Ve a Historial → Busca la actividad → Haz clic en la foto → Opciones → Eliminar."),
    ("¿Se pueden tomar notas de voz?", "Sí, tanto en Chat como en notas de audio. La app graba en formato MP3 y permite reproducción a velocidad variable."),
    ("¿Qué privacidad tienen mis datos?", "Todos los datos se encriptan en tránsito y en reposo. La geolocalización se almacena solo 7 días. Las fotos se borran después de 30 días si no están vinculadas a un documento oficial."),
    ("¿Puedo usar la app en varios dispositivos simultáneamente?", "Un usuario puede estar activo en máximo 2 dispositivos. Si intentas loguearte en un tercero, se desconectará automáticamente del primero."),
    ("¿Qué hacer si la app falla o crashea?", "Intenta: 1) Reinicia la app. 2) Limpia el caché de la app. 3) Reinicia el dispositivo. 4) Desinstala y reinstala la app. Si persiste, contacta al soporte."),
]

for pregunta, respuesta in faqs:
    p = doc.add_paragraph()
    p_run = p.add_run(f"P: {pregunta}")
    p_run.font.bold = True
    p_run.font.name = 'Arial'
    p_run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    
    r = doc.add_paragraph(f"R: {respuesta}")
    for run in r.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    r.paragraph_format.line_spacing = 1.5
    r.paragraph_format.left_indent = Inches(0.25)

doc.add_page_break()

# ============ 8. TROUBLESHOOTING ============
add_heading(doc, "8. TROUBLESHOOTING", 1, spacing_before=0, spacing_after=12)

add_heading(doc, "8.1. Problemas Comunes y Soluciones", 2, spacing_before=6, spacing_after=6)

troubles = [
    ("La cámara no funciona", "- Verifica que has dado permisos de cámara a la app (Configuración → CuidaLink → Cámara).\n- Reinicia la app e intenta de nuevo.\n- Comprueba que no hay otra app usando la cámara."),
    ("El GPS no funciona", "- Activa el GPS en Configuración del dispositivo.\n- Intenta hacer zoom out en el mapa para que se centre en tu ubicación.\n- Espera 30 segundos para que se calibre el GPS.\n- Si persiste, reinicia el dispositivo."),
    ("No recibo notificaciones", "- Verifica Configuración → Notificaciones → CuidaLink y asegúrate que está activado.\n- Comprueba que el sonido del dispositivo no está silenciado.\n- Intenta desactivar y reactivar notificaciones en la app."),
    ("La sincronización es lenta", "- Comprueba tu conexión a internet (Wi-Fi o datos móviles).\n- Intenta cerrar la app completamente y reabrirla.\n- Si la conexión es 3G o lenta, puede tardar más tiempo."),
    ("Olvidé mi PIN/Contraseña", "- Haz clic en 'Olvidé mi PIN' y sigue el proceso de recuperación.\n- Si no recibes el email, comprueba la carpeta de SPAM.\n- Contacta al administrador del sistema."),
    ("La app continúa mostrando 'no conectado'", "- Verifica que tienes conexión a internet.\n- Intenta cambiar entre Wi-Fi y datos móviles.\n- Reinicia el router/móvil.\n- Comprueba si hay mantenimiento del servidor (consulta el estado en la web)."),
]

for problema, solucion in troubles:
    p = doc.add_paragraph()
    p_run = p.add_run(f"Problema: {problema}")
    p_run.font.bold = True
    p_run.font.name = 'Arial'
    p_run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    
    s = doc.add_paragraph(f"Solución:\n{solucion}")
    for run in s.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    s.paragraph_format.line_spacing = 1.5
    s.paragraph_format.left_indent = Inches(0.25)

add_heading(doc, "8.2. Borrado de Caché y Datos", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Si la aplicación se comporta de manera extraña, puedes borrar el caché sin perder tus datos:",
alignment=WD_ALIGN_PARAGRAPH.LEFT)

cache_steps = [
    "ANDROID: Configuración → Aplicaciones → CuidaLink → Almacenamiento → Borrar Caché.",
    "iOS: Configuración → General → iPhone Storage → Selecciona CuidaLink → Descargar app.",
    "Abre la app de nuevo. Debería funcionar normalmente.",
]

for step in cache_steps:
    cs = doc.add_paragraph(step)
    for run in cs.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    cs.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ============ 9. CONTACTO Y SOPORTE ============
add_heading(doc, "9. CONTACTO Y SOPORTE", 1, spacing_before=0, spacing_after=12)

add_paragraph_text(doc,
"Si encontras problemas o tienes preguntas, no dudes en contactarnos:")

contact_info = [
    ("Email de Soporte", "soporte@cuidalink.es"),
    ("Teléfono", "+34 912 345 678"),
    ("Horario de Atención", "Lunes a Viernes, 9:00 - 18:00"),
    ("Chat en Vivo", "Disponible en la app (pestaña Ayuda)"),
    ("Base de Conocimientos", "https://help.cuidalink.es"),
    ("Reportar un Bug", "bugs@cuidalink.es"),
]

for tipo, info in contact_info:
    c = doc.add_paragraph(f"{tipo}: {info}", style='List Bullet')
    for run in c.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    c.paragraph_format.line_spacing = 1.5

add_paragraph_text(doc,
"\n¡Gracias por usar CuidaLink! Tu feedback nos ayuda a mejorar continuamente la plataforma.",
font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)

# Save document
output_path = r"c:\Users\usuario\Desktop\TFG ALZHEIMER\MANUAL_USUARIO_CUIDALINK.docx"
doc.save(output_path)
print(f"✓ Manual de usuario creado: {output_path}")
print(f"✓ Formato: DIN A4, Arial 12pt, espaciado 1.5")
print(f"✓ Secciones: Introducción, Requisitos, Instalación, Guías Cuidadora y Familiar, FAQ, Troubleshooting, Contacto")
