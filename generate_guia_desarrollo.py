#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1, spacing_before=12, spacing_after=6):
    """Add a styled heading"""
    heading = doc.add_paragraph()
    heading_run = heading.add_run(text)
    
    if level == 1:
        heading_run.font.size = Pt(18)
        heading_run.font.bold = True
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
    else:
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
    
    heading_run.font.name = 'Arial'
    heading.paragraph_format.line_spacing = 1.5
    heading.paragraph_format.space_before = Pt(spacing_before)
    heading.paragraph_format.space_after = Pt(spacing_after)
    
    return heading

def add_paragraph_text(doc, text, font_size=12, spacing_line=1.5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Add a paragraph with consistent formatting"""
    paragraph = doc.add_paragraph(text)
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(font_size)
    paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = spacing_line
    return paragraph

def add_code_block(doc, code, language=""):
    """Add a code block with monospace font"""
    code_para = doc.add_paragraph()
    code_para.paragraph_format.left_indent = Inches(0.5)
    code_para.paragraph_format.space_before = Pt(6)
    code_para.paragraph_format.space_after = Pt(6)
    
    # Add language label if provided
    if language:
        label = code_para.add_run(f"[{language}]\n")
        label.font.name = 'Courier New'
        label.font.size = Pt(10)
        label.font.italic = True
    
    code_run = code_para.add_run(code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)
    
    return code_para

# Create document
doc = Document()

# ============ PORTADA ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("GUÍA DE DESARROLLO")
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.name = 'Arial'
title.paragraph_format.space_after = Pt(6)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("CuidaLink")
subtitle_run.font.size = Pt(24)
subtitle_run.font.name = 'Arial'
subtitle_run.font.bold = True
subtitle.paragraph_format.space_after = Pt(12)

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc_run = desc.add_run("Para Desarrolladores\nCómo Entender y Trabajar con el Código")
desc_run.font.size = Pt(14)
desc_run.font.name = 'Arial'
desc_run.font.italic = True
desc.paragraph_format.space_after = Pt(24)

for _ in range(6):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info.add_run("Versión 1.0\nMarzo de 2026")
info_run.font.size = Pt(12)
info_run.font.name = 'Arial'

doc.add_page_break()

# ============ ÍNDICE ============
add_heading(doc, "ÍNDICE", 1, spacing_before=0, spacing_after=12)

toc_items = [
    "1. Descripción General del Proyecto",
    "2. Configuración del Entorno",
    "3. Estructura del Proyecto",
    "4. Frontend (React Native + Expo)",
    "   4.1. Estructura de Carpetas",
    "   4.2. Key Files y Componentes",
    "   4.3. Servicios Principales",
    "   4.4. Estado Global (Context API)",
    "   4.5. Cómo Agregar una Nueva Pantalla",
    "5. Backend (Spring Boot)",
    "   5.1. Estructura de Carpetas",
    "   5.2. Arquitectura: Controller → Service → Repository",
    "   5.3. Entidades de Base de Datos",
    "   5.4. Endpoints Principales",
    "   5.5. Autenticación JWT",
    "   5.6. Cómo Agregar un Nuevo Endpoint",
    "6. Base de Datos",
    "   6.1. Esquema de Tablas",
    "   6.2. Relaciones",
    "   6.3. Configuración Local vs Producción",
    "7. Flujos Principales",
    "   7.1. Autenticación",
    "   7.2. Registro de Evento (Tarea)",
    "   7.3. Chat en Tiempo Real",
    "8. Git Workflow",
    "9. Estándares de Código",
    "10. Testing",
    "11. Deployment",
    "12. Troubleshooting Técnico",
    "13. Recursos Útiles",
]

for item in toc_items:
    toc_para = doc.add_paragraph(item)
    toc_para.paragraph_format.left_indent = Inches(0.25) if item.startswith("   ") else Inches(0)
    for run in toc_para.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    toc_para.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ============ 1. DESCRIPCIÓN GENERAL ============
add_heading(doc, "1. DESCRIPCIÓN GENERAL DEL PROYECTO", 1, spacing_before=0, spacing_after=12)

add_paragraph_text(doc,
"CuidaLink es una plataforma de cuidado de personas con Alzheimer que conecta cuidadores profesionales con familias. La arquitectura es de tres capas: Frontend móvil (React Native), Backen API (Spring Boot), Base de datos en la nube (PostgreSQL).")

add_heading(doc, "1.1. Stack Tecnológico", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Frontend: React Native 0.81 + Expo SDK 54 + TypeScript 5.9\n"
"Backend: Spring Boot 3.2 + Java 17 + Spring Security + JWT\n"
"Base de Datos: PostgreSQL 15+ (Neon Cloud en producción, H2 en desarrollo)\n"
"Control de Versiones: Git (GitHub: jocabu05/CuidaLink)\n"
"Build Tools: Maven 3.8.5 (backend), npm/yarn (frontend)")

add_heading(doc, "1.2. Contexto del Proyecto", 2, spacing_before=6, spacing_after=6)
add_paragraph_text(doc,
"Este es un Trabajo Final de Grado (TFG) desarrollado en 2 semanas. El objetivo es crear una MVP (Minimum Viable Product) funcional que demuestre la viabilidad técnica y comercial de una plataforma de cuidado digital. El código está completamente funcional pero puede requerir optimizaciones para producción a escala.")

doc.add_page_break()

# ============ 2. CONFIGURACIÓN DEL ENTORNO ============
add_heading(doc, "2. CONFIGURACIÓN DEL ENTORNO", 1, spacing_before=0, spacing_after=12)

add_heading(doc, "2.1. Requisitos Previos", 2, spacing_before=6, spacing_after=6)

reqs = [
    ("Node.js", "Versión 18+", "Descargable de nodejs.org"),
    ("Java", "JDK 17+", "Descargable de oracle.com o openjdk.java.net"),
    ("Maven", "3.8+", "Incluido en el wrapper del proyecto"),
    ("Git", "Última versión", "Descargable de git-scm.com"),
    ("VS Code", "Recomendado", "Descargable de code.visualstudio.com"),
    ("Android Emulator/Xcode", "Para testing móvil", "Android Studio o Xcode según tu OS"),
]

for nom, version, fuente in reqs:
    p = doc.add_paragraph(f"{nom}: {version} ({fuente})", style='List Bullet')
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5

add_heading(doc, "2.2. Clonar el Repositorio", 2, spacing_before=6, spacing_after=6)

add_code_block(doc, "git clone https://github.com/jocabu05/CuidaLink.git\ncd CuidaLink", "bash")

add_heading(doc, "2.3. Configurar Backend", 2, spacing_before=6, spacing_after=6)

add_paragraph_text(doc, "Desarrollo Local (H2 in-memory):")
add_code_block(doc, "cd backend\nmvn clean install\nmvn spring-boot:run\n# Arranca en http://localhost:8080", "bash")

add_paragraph_text(doc, "Verificar que funciona:")
add_code_block(doc, "curl http://localhost:8080/api/auth/login \\  \n  -X POST \\  \n  -H 'Content-Type: application/json' \\  \n  -d '{\"telefono\":\"645123456\",\"pin\":\"1234\"}'", "bash")

add_heading(doc, "2.4. Configurar Frontend", 2, spacing_before=6, spacing_after=6)

add_code_block(doc, "cd mobile\nnpm install\n# En desarrollo con Expo\nnpx expo start\n# Escanea QR con Expo Go o pulsa 'a' para Android", "bash")

add_paragraph_text(doc, "Para build de producción:")
add_code_block(doc, "eas build --platform android  # APK para Android\neas build --platform ios     # IPA para iOS", "bash")

doc.add_page_break()

# ============ 3. ESTRUCTURA DEL PROYECTO ============
add_heading(doc, "3. ESTRUCTURA DEL PROYECTO", 1, spacing_before=0, spacing_after=12)

add_code_block(doc, """CuidaLink/
├── backend/                      # API REST Spring Boot
│   ├── pom.xml                   # Dependencias Maven
│   ├── src/main/java/com/cuidalink/
│   │   ├── config/               # Configuración (CORS, DataInitializer)
│   │   ├── controller/           # REST Controllers (6 controllers)
│   │   ├── dto/                  # Data Transfer Objects
│   │   ├── entity/               # JPA Entities (9 entidades)
│   │   ├── repository/           # Spring Data JPA Repositories
│   │   ├── security/             # JWT, UserDetailsService, Filters
│   │   ├── service/              # Lógica de negocio (6 services)
│   │   └── CuidaLinkApplication.java
│   └── src/main/resources/
│       ├── application.yml       # Config local (H2)
│       ├── application-prod.yml  # Config producción (PostgreSQL/Neon)
│       └── data-h2.sql          # Seed data para H2
│
├── mobile/                       # App React Native + Expo
│   ├── package.json
│   ├── App.tsx                   # Navegación principal (Navigator)
│   ├── src/
│   │   ├── screens/              # 23 pantallas (4 helpers, 19 screens)
│   │   ├── components/           # 4 componentes reutilizables
│   │   ├── context/              # ThemeContext (estado global)
│   │   ├── services/             # 18 servicios (API, GPS, chat, etc)
│   │   ├── styles/               # Tema unificado (colores, tipografía)
│   │   └── __tests__/            # Tests unitarios (Jest)
│   └── assets/                   # Imágenes, iconos, fuentes
│
├── database/                     # Esquemas SQL
│   ├── schema.sql                # Esquema completo (H2)
│   ├── schema-postgres.sql       # Esquema PostgreSQL
│   └── data-postgres.sql         # Seed data
│
└── README.md                     # Documentación proyecto""", "text")

doc.add_page_break()

# ============ 4. FRONTEND ============
add_heading(doc, "4. FRONTEND (REACT NATIVE + EXPO)", 1, spacing_before=0, spacing_after=12)

add_heading(doc, "4.1. Estructura de Carpetas", 2, spacing_before=6, spacing_after=6)

add_paragraph_text(doc, "mobile/src/ contiene:")

front_folders = [
    ("screens/", "23 archivos .tsx con las pantallas de la app"),
    ("components/", "BigButton, TaskCard, WearableWidget, TaskEditorModal"),
    ("services/", "18 servicios: authService, apiService, locationService, chatService, etc."),
    ("context/", "ThemeContext para estado global (tema oscuro, rol, usuario actual)"),
    ("styles/", "Archivo theme.ts con colores, tipografía, espacios estandarizados"),
    ("__tests__/", "Tests unitarios con Jest para servicios críticos"),
]

for folder, desc in front_folders:
    fp = doc.add_paragraph(f"{folder}: {desc}", style='List Bullet')
    for run in fp.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    fp.paragraph_format.line_spacing = 1.5

add_heading(doc, "4.2. Key Files y Componentes", 2, spacing_before=6, spacing_after=6)

key_files = [
    ("App.tsx", "Punto de entrada. Configura Stack Navigator, Tabs, ThemeContext wrapper"),
    ("src/screens/DashboardScreen.tsx", "Pantalla principal cuidadora (919 líneas). Muestra tareas, notas, barra de progreso"),
    ("src/screens/FamiliarDashboardScreen.tsx", "Pantalla principal familiar (963 líneas). Panel de supervisión"),
    ("src/screens/MapaZonaScreen.tsx", "Mapa interactivo con Leaflet (WebView). ~50 POIs hardcoded, zona segura circular"),
    ("src/screens/PaseoScreen.tsx", "Tracking GPS en tiempo real. Registra ubicación cada 10 seg, calcula distancia"),
    ("src/services/authService.ts", "Manejo de JWT, login/logout, persistencia con AsyncStorage"),
    ("src/services/apiService.ts", "Cliente Axios configurado. Interceptores para JWT, manejo de errores, retry"),
    ("src/services/locationService.ts", "Geocodificación, geofencing, cálculo de distancias"),
    ("src/services/chatService.ts", "CRUD de mensajes, sincronización local-servidor"),
]

for file, desc in key_files:
    kf = doc.add_paragraph(f"{file}: {desc}", style='List Bullet')
    for run in kf.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    kf.paragraph_format.line_spacing = 1.5

add_heading(doc, "4.3. Servicios Principales", 2, spacing_before=6, spacing_after=6)

add_paragraph_text(doc, "authService.ts - Autenticación con JWT:")
add_code_block(doc, "// Login: envía teléfono + PIN al backend, recibe JWT\nawait authService.login('645123456', '1234');\n// Obtener usuario actual\nconst user = authService.getCurrentUser();\n// Logout: borra JWT de AsyncStorage\nauthService.logout();", "typescript")

add_paragraph_text(doc, "eventosService.ts - Registro de eventos:")
add_code_block(doc, "// Crear evento (llegada, paseo, medicación, etc.)\nawait eventosService.createEvento({\n  tipo: 'LLEGADA',\n  timestamp: new Date(),\n  fotoUrl: 'base64-image',\n  ubicacion: { lat: 39.1525, lng: -0.4345 }\n});", "typescript")

add_paragraph_text(doc, "locationService.ts - Geolocalización:")
add_code_block(doc, "// Obtener ubicación actual (precision ~10-20m)\nconst location = await locationService.getCurrentLocation();\n// Calcular distancia entre dos puntos\nconst distance = locationService.calculateDistance(lat1, lng1, lat2, lng2);\n// Detectar si está fuera de zona segura\nconst isOutOfZone = locationService.isOutsideSafeZone(location, safeZone);", "typescript")

add_heading(doc, "4.4. Estado Global (Context API)", 2, spacing_before=6, spacing_after=6)

add_paragraph_text(doc, "ThemeContext proporciona:")
add_code_block(doc, "// En cualquier screen, acceder al contexto\nconst { theme, isDarkMode, userRole, currentUser } = useContext(ThemeContext);\n\n// Cambiar tema\ntheme.toggleDarkMode();\n\n// Verificar rol\nif (userRole === 'CUIDADORA') {\n  // Mostrar pantalla de cuidadora\n}", "typescript")

add_heading(doc, "4.5. Cómo Agregar una Nueva Pantalla", 2, spacing_before=6, spacing_after=6)

add_paragraph_text(doc, "Ejemplo: Agregar pantalla 'ListadoCumplidoScreen':")

steps = [
    "Crear archivo: mobile/src/screens/ListadoCumplidoScreen.tsx",
    "Importar componentes necesarios (View, Text, FlatList, etc.) y servicios",
    "Crear component React Native con hooks (useState, useEffect, useContext)",
    "Agregar navegación en App.tsx en la sección de Stack.Screen",
    "Usar servicios para obtener datos: tareasService.getTareasCumplidas()",
    "Aplicar estilos del theme.ts (colores, espacios, tipografía)",
    "Si hay lógica compleja, crear servicio específico en services/",
    "Agregar tests unitarios en __tests__/",
]

for i, step in enumerate(steps, 1):
    s = doc.add_paragraph(f"{i}. {step}")
    for run in s.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    s.paragraph_format.line_spacing = 1.5

add_code_block(doc, """// mobile/src/screens/ListadoCumplidoScreen.tsx
import React, { useState, useEffect, useContext } from 'react';
import { View, FlatList, Text, StyleSheet } from 'react-native';
import { ThemeContext } from '../context/ThemeContext';
import { tareasService } from '../services/tareasService';

export const ListadoCumplidoScreen = ({ navigation }) => {
  const { theme } = useContext(ThemeContext);
  const [tareas, setTareas] = useState([]);

  useEffect(() => {
    cargarTareasCumplidas();
  }, []);

  const cargarTareasCumplidas = async () => {
    try {
      const tareasObtenidas = await tareasService.getTareasCumplidas();
      setTareas(tareasObtenidas);
    } catch (error) {
      console.error('Error cargando tareas:', error);
    }
  };

  return (
    <View style={[styles.container, { backgroundColor: theme.backgroundColor }]}>
      <FlatList
        data={tareas}
        keyExtractor={(item) => item.id}
        renderItem={({ item }) => (
          <Text style={[styles.item, { color: theme.textColor }]}>{item.nombre}</Text>
        )}
      />
    </View>
  );
};

const styles = StyleSheet.create({
  container: { flex: 1, padding: 16 },
  item: { fontSize: 16, marginVertical: 8 },
});""", "typescript")

doc.add_page_break()

# ============ 5. BACKEND ============
add_heading(doc, "5. BACKEND (SPRING BOOT)", 1, spacing_before=0, spacing_after=12)

add_heading(doc, "5.1. Estructura de Carpetas", 2, spacing_before=6, spacing_after=6)

add_paragraph_text(doc, "backend/src/main/java/com/cuidalink/:")

back_folders = [
    ("config/", "GlobalExceptionHandler, CORS config, DataInitializer (seed data)"),
    ("controller/", "6 REST controllers para autenticación, eventos, chat, etc."),
    ("dto/", "Data Transfer Objects para request/response"),
    ("entity/", "9 entidades JPA (Abuelo, Cuidadora, Familiar, Evento, etc.)"),
    ("repository/", "Spring Data JPA repositories (queries personalizadas)"),
    ("security/", "JWT utils, UserDetailsService, security config"),
    ("service/", "Lógica de negocio (EventoService, CuidadoraService, etc.)"),
]

for folder, desc in back_folders:
    bf = doc.add_paragraph(f"{folder}: {desc}", style='List Bullet')
    for run in bf.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    bf.paragraph_format.line_spacing = 1.5

add_heading(doc, "5.2. Arquitectura: Controller → Service → Repository", 2, spacing_before=6, spacing_after=6)

add_code_block(doc, """// 1. CONTROLLER - Maneja HTTP requests/responses
@RestController
@RequestMapping(\"/api/eventos\")
public class EventoController {
  @PostMapping
  public ResponseEntity<?> crearEvento(@RequestBody EventoDTO dto) {
    return eventoService.crearEvento(dto);
  }
}

// 2. SERVICE - Lógica de negocio
@Service
public class EventoService {
  public Evento crearEvento(EventoDTO dto) {
    Evento evento = new Evento();
    evento.setTipo(dto.getTipo());
    evento.setTimestamp(new Date());
    return eventoRepository.save(evento);
  }
}

// 3. REPOSITORY - Acceso a datos (JPA)
@Repository
public interface EventoRepository extends JpaRepository<Evento, Long> {
  List<Evento> findByCuidadoraIdOrderByTimestampDesc(Long cuidadoraId);
}""", "java")

add_heading(doc, "5.3. Entidades de Base de Datos", 2, spacing_before=6, spacing_after=6)

entities_desc = [
    "Abuelo: Persona con Alzheimer (nombre, edad, dirección, medicamentos activos)",
    "Cuidadora: Profesional (nombre, teléfono, PIN hash, rating promedio)",
    "Familiar: Miembro familia (nombre, email, contraseña, relación)",
    "Evento: Actividad registrada (tipo: LLEGADA/PASTILLA/COMIDA/PASEO, timestamp, foto)",
    "Cita: Cita médica (fecha, hora, descripción, recordatorio)",
    "Nota: Mensaje familia→cuidadora (contenido, urgente, leída)",
    "MensajeChat: Historial chat (remitente, contenido, timestamp)",
    "Paseo: Registro paseo (inicio, fin, distancia, ruta JSON)",
    "Rating: Valoración (puntuación 1-5, comentario, timestamp)",
]

for ent in entities_desc:
    ed = doc.add_paragraph(ent, style='List Bullet')
    for run in ed.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    ed.paragraph_format.line_spacing = 1.5

add_heading(doc, "5.4. Endpoints Principales", 2, spacing_before=6, spacing_after=6)

endpoints = [
    ("POST /api/auth/login", "{ telefono, pin } → { token, usuario }"),
    ("POST /api/auth/refresh", "Refrescar JWT expirado"),
    ("POST /api/eventos", "Crear nuevo evento"),
    ("GET /api/eventos/{id}", "Obtener detalles evento"),
    ("GET /api/cuidadoras/{id}/eventos", "Listar eventos por cuidadora"),
    ("POST /api/chat/mensajes", "Enviar mensaje chat"),
    ("GET /api/chat/mensajes/{conversationId}", "Obtener historial"),
    ("POST /api/notas", "Enviar nota del familiar"),
    ("GET /api/notas/{cuidadoraId}", "Obtener notas por cuidadora"),
]

for endpoint, desc in endpoints:
    ep = doc.add_paragraph(f"{endpoint}: {desc}", style='List Bullet')
    for run in ep.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    ep.paragraph_format.line_spacing = 1.5

add_heading(doc, "5.5. Autenticación JWT", 2, spacing_before=6, spacing_after=6)

add_code_block(doc, """// 1. Login: validar teléfono+PIN, generar JWT
POST /api/auth/login
{ 
  \"telefono\": \"645123456\", 
  \"pin\": \"1234\" 
}

// Response
{
  \"token\": \"eyJhbGc...\",
  \"cuidadoraId\": 1,
  \"nombre\": \"María García\"
}

// 2. Usar token en headers
Authorization: Bearer eyJhbGc...

// 3. Validar token en backend (interceptor)
@Component
public class JwtFilter extends OncePerRequestFilter {
  protected void doFilterInternal(HttpServletRequest req...) {
    String token = extractToken(req);
    if (jwtUtils.validateToken(token)) {
      Long userId = jwtUtils.getUserIdFromToken(token);
      // Autenticar usuario
    }
  }
}""", "java")

add_heading(doc, "5.6. Cómo Agregar un Nuevo Endpoint", 2, spacing_before=6, spacing_after=6)

add_paragraph_text(doc, "Ejemplo: Agregar endpoint para listar medicamentos:")

endpoint_steps = [
    "Crear DTO: medicamento/MedicamentoDTO.java (mapeador de datos)",
    "Crear Entity: entity/Medicamento.java con @Entity, @Id, getters/setters",
    "Crear Repository: repository/MedicamentoRepository.java (extends JpaRepository)",
    "Crear Service: service/MedicamentoService.java con lógica de negocio",
    "Crear Controller: controller/MedicamentoController.java con @RestController",
    "Definir endpoints: @GetMapping, @PostMapping, @PutMapping, @DeleteMapping",
    "Agregar tests en __tests__/",
    "Documentar en archivo README o ARQUITECTURA",
]

for i, step in enumerate(endpoint_steps, 1):
    es = doc.add_paragraph(f"{i}. {step}")
    for run in es.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    es.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ============ 6. BASE DE DATOS ============
add_heading(doc, "6. BASE DE DATOS", 1, spacing_before=0, spacing_after=12)

add_heading(doc, "6.1. Esquema de Tablas", 2, spacing_before=6, spacing_after=6)

add_code_block(doc, """-- 9 tablas principales
ABUELOS (id, nombre, edad, dni, direccion, lat, lng)
CUIDADORAS (id, nombre, telefono, pin_hash, rating)
FAMILIARES (id, nombre, email, contraseña_hash, relacion)
EVENTOS (id, cuidadora_id, abuelo_id, tipo, timestamp, foto_url, ubicacion)
CITAS (id, abuelo_id, fecha, hora, descripcion, recordatorio)
NOTAS (id, familiar_id, cuidadora_id, contenido, urgente, leida)
MENSAJES_CHAT (id, remitente_id, destinatario_id, contenido, timestamp, tipo_medio)
PASEOS (id, cuidadora_id, abuelo_id, inicio, fin, distancia, ruta_json)
RATINGS (id, cuidadora_id, puntuacion, comentario, timestamp)""", "sql")

add_heading(doc, "6.2. Relaciones", 2, spacing_before=6, spacing_after=6)

relations = [
    "CUIDADORA (1) → (N) EVENTO",
    "ABUELO (1) → (N) EVENTO",
    "FAMILIAR (1) → (N) NOTA",
    "CUIDADORA (1) → (N) NOTA (receptora)",
    "CUIDADORA (1) → (N) PASEO",
    "CUIDADORA (1) → (N) RATING",
]

for rel in relations:
    r = doc.add_paragraph(rel, style='List Bullet')
    for run in r.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    r.paragraph_format.line_spacing = 1.5

add_heading(doc, "6.3. Configuración Local vs Producción", 2, spacing_before=6, spacing_after=6)

add_paragraph_text(doc, "DESARROLLO (application.yml - H2 in-memory):")
add_code_block(doc, """spring:
  datasource:
    url: jdbc:h2:mem:testdb
    driver-class-name: org.h2.Driver
  jpa:
    hibernate:
      ddl-auto: create-drop
    database-platform: org.hibernate.dialect.H2Dialect""", "yaml")

add_paragraph_text(doc, "PRODUCCIÓN (application-prod.yml - PostgreSQL/Neon):")
add_code_block(doc, """spring:
  datasource:
    url: jdbc:postgresql://neon-host/neondb?sslmode=require
    username: neondb_owner
    password: ****
  jpa:
    hibernate:
      ddl-auto: update
    database-platform: org.hibernate.dialect.PostgreSQL10Dialect""", "yaml")

doc.add_page_break()

# ============ 7. FLUJOS PRINCIPALES ============
add_heading(doc, "7. FLUJOS PRINCIPALES", 1, spacing_before=0, spacing_after=12)

add_heading(doc, "7.1. Flujo de Autenticación", 2, spacing_before=6, spacing_after=6)

add_code_block(doc, """1. Mobile: Usuario introduce teléfono + PIN
2. Mobile → Backend: POST /api/auth/login { telefono, pin }
3. Backend: Buscar Cuidadora en BD
4. Backend: Validar PIN con BCrypt
5. Backend: Generar JWT token (válido 24h)
6. Backend → Mobile: ResponseEntity { token, usuarioId, nombre }
7. Mobile: Guarda token en AsyncStorage (encriptado)
8. Mobile: Redirige a DashboardScreen
9. Endpoints posteriores incluyen: Authorization: Bearer {token}""", "text")

add_heading(doc, "7.2. Flujo de Registro de Evento (Tarea)", 2, spacing_before=6, spacing_after=6)

add_code_block(doc, """1. Cuidadora hace clic en 'Pastilla' → abre cámara
2. Cuidadora fotografía medicamento
3. Mobile: Captura foto, convierte a Base64
4. Mobile: Envía POST /api/eventos
   {
     \"tipo\": \"PASTILLA\",
     \"timestamp\": \"2026-03-04T14:30:00Z\",
     \"fotoUrl\": \"data:image/jpeg;base64,...\",
     \"ubicacion\": { \"lat\": 39.15, \"lng\": -0.43 }
   }
5. Backend: EventoService.crearEvento()
6. Backend: Guarda foto en storage (local o cloud)
7. Backend: Inserta evento en BD
8. Backend → Mobile: { eventoId, success: true }
9. Mobile: Marca tarea como completada en UI
10. Familia: Recibe notificación push de evento
11. Familia: Ve evento en timeline del panel""", "text")

add_heading(doc, "7.3. Flujo de Chat en Tiempo Real", 2, spacing_before=6, spacing_after=6)

add_code_block(doc, """1. Cuidadora escribe mensaje en ChatScreen
2. Mobile (Cuidadora): POST /api/chat/mensajes
   {
     \"remitente\": \"cuidadora_id\",
     \"destinatario\": \"familiar_id\",
     \"contenido\": \"Medicamento administrado\",
     \"tipoMedio\": \"TEXT\"
   }
3. Backend: Guarda mensaje en BD
4. Backend: Envía notificación push a Familiar (WebSocket/Firebase)
5. Mobile (Familiar): Recibe notificación
6. Mobile (Familiar): Se actualiza ChatScreen (FlatList)
7. Backend: Actualiza timestamp de lectura cuando Familiar abre chat
8. Mobile (Cuidadora): Ve ✓✓ (leído) en su mensaje
9. Historial persiste en BD para acceso offline (caché local)""", "text")

doc.add_page_break()

# ============ 8. GIT WORKFLOW ============
add_heading(doc, "8. GIT WORKFLOW", 1, spacing_before=0, spacing_after=12)

add_paragraph_text(doc, "El proyecto usa Git con rama main para producción. Para agregar nuevas funcionalidades:")

git_steps = [
    "Crear rama desde main: git checkout -b feature/nombre-funcionalidad",
    "Realizar cambios en local: backend/ o mobile/",
    "Hacer commits frecuentes: git commit -m 'feat: descripción del cambio'",
    "Subir rama: git push origin feature/nombre-funcionalidad",
    "Crear Pull Request en GitHub",
    "Revisión de código (en equipo real)",
    "Merge a main: git merge feature/nombre-funcionalidad",
    "Delete rama: git branch -d feature/nombre-funcionalidad",
]

for i, step in enumerate(git_steps, 1):
    gs = doc.add_paragraph(f"{i}. {step}")
    for run in gs.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    gs.paragraph_format.line_spacing = 1.5

add_heading(doc, "8.1. Mensajes de Commit", 2, spacing_before=6, spacing_after=6)

add_paragraph_text(doc, "Usar prefijos estándar:")
add_code_block(doc, """feat: nueva funcionalidad (Agregar chat de audio)\nfix: corrección de bug (Validar PIN con mayúsculas)\nrefactor: reestructuración sin cambios funcionales\ndocs: cambios de documentación\ntest: agregar o actualizar tests\nstyle: formateo de código (no cambia lógica)""", "text")

doc.add_page_break()

# ============ 9. ESTÁNDARES DE CÓDIGO ============
add_heading(doc, "9. ESTÁNDARES DE CÓDIGO", 1, spacing_before=0, spacing_after=12)

add_heading(doc, "9.1. TypeScript/JavaScript (Frontend)", 2, spacing_before=6, spacing_after=6)

add_code_block(doc, """// Nombres de archivos
- PascalCase para screens: DashboardScreen.tsx
- camelCase para services: eventosService.ts
- kebab-case para eventos: on-task-complete.ts

// Tipos
interface User {
  id: number;
  nombre: string;
  email: string;
}

// Funciones flecha
const getUserData = async (id: number): Promise<User> => {
  return await apiService.get(`/users/${id}`);
};

// Destructuring
const { nombre, email } = usuario;

// Constants
const MAX_RETRY_ATTEMPTS = 3;
const DEFAULT_TIMEOUT = 5000;""", "typescript")

add_heading(doc, "9.2. Java (Backend)", 2, spacing_before=6, spacing_after=6)

add_code_block(doc, """// Nombres de clases
- PascalCase para todas: CuidadoraService, EventoController, Usuario

// Métodos
- camelCase: crearEvento(), obtenerEventosPorCuidadora()
- Verbos: get, set, create, update, delete, find

// Constantes
private static final String DEFAULT_ENCODING = \"UTF-8\";
private static final int TIMEOUT_SECONDS = 30;

// Anotaciones
@Service
public class EventoService {
  @Autowired
  private EventoRepository eventoRepository;
  
  @Transactional
  public Evento crearEvento(EventoDTO dto) { ... }
}

// Manejo de excepciones
try {
  // lógica
} catch (DataAccessException e) {
  log.error(\"Error accediendo BD\", e);
  throw new RuntimeException(\"Error accediendo datos\", e);
}""", "java")

add_heading(doc, "9.3. Convenciones Generales", 2, spacing_before=6, spacing_after=6)

conventions = [
    "Comentarios: Solo para lógica compleja. Code debe ser auto-documentable.",
    "Máximo 80 caracteres por línea (excepción: URLs, logs)",
    "Indentación: 2 espacios en TypeScript, 4 en Java",
    "Imports: Ordenados alfabéticamente, agrupar por tipo",
    "Archivos vacíos o sin uso: Eliminar",
    "Passwords/Tokens hardcodeados: NUNCA. Usar variables de entorno",
    "Logs: Usar niveles (DEBUG, INFO, WARN, ERROR) según criticidad",
]

for conv in conventions:
    c = doc.add_paragraph(conv, style='List Bullet')
    for run in c.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    c.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ============ 10. TESTING ============
add_heading(doc, "10. TESTING", 1, spacing_before=0, spacing_after=12)

add_heading(doc, "10.1. Frontend Tests (Jest)", 2, spacing_before=6, spacing_after=6)

add_paragraph_text(doc, "Tests unitarios para servicios críticos:")
add_code_block(doc, """// mobile/__tests__/reminderService.test.ts
describe('reminderService', () => {
  it('debería programar recordatorio correctamente', async () => {
    const resultado = await reminderService.programarRecordatorio({
      titulo: 'Pastilla',
      hora: '09:00'
    });
    expect(resultado.success).toBe(true);
  });

  it('debería fallar sin hora válida', async () => {
    expect(() => reminderService.obtenerMinutosHasta('hora-invalida'))
      .toThrow('Hora inválida');
  });
});

// Ejecutar tests
npm test""", "typescript")

add_heading(doc, "10.2. Backend Tests (JUnit)", 2, spacing_before=6, spacing_after=6)

add_code_block(doc, """// backend/src/test/java/...
@SpringBootTest
public class EventoControllerTest {
  @Autowired
  private MockMvc mockMvc;
  
  @Test
  public void crearEventoDeberiaRetornar201() throws Exception {
    mockMvc.perform(post(\"/api/eventos\")
      .contentType(MediaType.APPLICATION_JSON)
      .content(\"{...}\"))
      .andExpect(status().isCreated());
  }
}

// Ejecutar tests
mvn test""", "java")

add_heading(doc, "10.3. Cobertura de Código", 2, spacing_before=6, spacing_after=6)

add_paragraph_text(doc, "Objetivo: Mínimo 75% cobertura en servicios críticos.")
add_code_block(doc, """# Frontend
npm test -- --coverage

# Backend
mvn test jacoco:report
# Reporte en target/site/jacoco/index.html""", "bash")

doc.add_page_break()

# ============ 11. DEPLOYMENT ============
add_heading(doc, "11. DEPLOYMENT", 1, spacing_before=0, spacing_after=12)

add_heading(doc, "11.1. Build Backend para Producción", 2, spacing_before=6, spacing_after=6)

add_code_block(doc, """# Compilar JAR
mvn clean package -DskipTests

# Archivo generado: backend/target/cuidalink-0.0.1-SNAPSHOT.jar

# Ejecutar en producción
java -jar cuidalink-0.0.1-SNAPSHOT.jar \\
  --spring.profiles.active=prod \\
  --server.port=8080""", "bash")

add_heading(doc, "11.2. Build Mobile para Producción", 2, spacing_before=6, spacing_after=6)

add_code_block(doc, """# Android APK
eas build --platform android --auto-submit

# iOS IPA (requiere Apple Developer Account)
eas build --platform ios

# Outputs
# Android: .../app.apk (45MB)
# iOS: .../app.ipa (60MB)""", "bash")

add_heading(doc, "11.3. Variables de Entorno", 2, spacing_before=6, spacing_after=6)

add_paragraph_text(doc, "Crear archivo .env en raíz del backend:")
add_code_block(doc, """DATABASE_URL=jdbc:postgresql://neon-host:5432/neondb
DATABASE_USER=neondb_owner
DATABASE_PASSWORD=your-password-here
JWT_SECRET=your-very-secret-key-123456789
CORS_ALLOWED_ORIGINS=https://yourdomain.com,https://app.yourdomain.com
STORAGE_PATH=/var/cuidalink/storage""", "env")

doc.add_page_break()

# ============ 12. TROUBLESHOOTING ============
add_heading(doc, "12. TROUBLESHOOTING TÉCNICO", 1, spacing_before=0, spacing_after=12)

troubleshooting_items = [
    ("npm install falla", "Borrar node_modules y package-lock.json. Reintentar: npm install"),
    ("Port 8080 ya en uso", "Cambiar puerto: mvn spring-boot:run -Dspring-boot.run.arguments=--server.port=8081"),
    ("Base de datos no conecta", "Verificar application.yml, comprobar permisos PostgreSQL"),
    ("JWT token inválido", "Verificar JWT_SECRET coincide, timestamp del servidor correcta"),
    ("Imágenes no cargan en chat", "Comprobar Storage Path existe, permisos lectura/escritura"),
    ("Push notifications no llegan", "Verificar Firebase config, permisos notificaciones en app"),
    ("App lenta en producción", "Optimizar queries (N+1 problem), agregar índices BD, caché"),
]

for problema, solucion in troubleshooting_items:
    p = doc.add_paragraph()
    p_run = p.add_run(f"{problema}: ")
    p_run.font.bold = True
    p.add_run(solucion)
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    p.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ============ 13. RECURSOS ÚTILES ============
add_heading(doc, "13. RECURSOS ÚTILES", 1, spacing_before=0, spacing_after=12)

add_heading(doc, "13.1. Documentación Oficial", 2, spacing_before=6, spacing_after=6)

resources = [
    ("React Native Docs", "https://reactnative.dev/docs"),
    ("Expo Documentation", "https://docs.expo.dev"),
    ("Spring Boot Docs", "https://spring.io/projects/spring-boot"),
    ("PostgreSQL Docs", "https://www.postgresql.org/docs"),
    ("Neon (PostgreSQL Cloud)", "https://neon.tech/docs"),
    ("JWT.io", "https://jwt.io - Debugger y generador JWT"),
    ("MDN Web Docs", "https://developer.mozilla.org"),
]

for recurso, url in resources:
    r = doc.add_paragraph(f"{recurso}: {url}")
    for run in r.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    r.paragraph_format.line_spacing = 1.5

add_heading(doc, "13.2. Herramientas Recomendadas", 2, spacing_before=6, spacing_after=6)

tools = [
    "VS Code + Extensions: React Native Tools, Spring Boot Extension Pack, Prettier",
    "Postman/Thunder Client: Para testear endpoints REST",
    "DBeaver: GUI para administrar bases de datos PostgreSQL",
    "Git Kraken: Cliente Git visual (opcional)",
    "Expo Go: App para preview móvil en desarrollo",
]

for tool in tools:
    t = doc.add_paragraph(tool, style='List Bullet')
    for run in t.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    t.paragraph_format.line_spacing = 1.5

add_heading(doc, "13.3. Contacto del Equipo", 2, spacing_before=6, spacing_after=6)

add_paragraph_text(doc,
"Autor: Jorge Castera Bueno\n"
"Email: jorgecb@example.com\n"
"GitHub: @jocabu05\n"
"Soporte técnico: soporte@cuidalink.es\n\n"
"¡Bienvenido al equipo de CuidaLink! Si tienes dudas, no temas preguntar. El código es para entenderlo, mejorar y evolucionar juntos.",
alignment=WD_ALIGN_PARAGRAPH.LEFT)

# Save document
output_path = r"c:\Users\usuario\Desktop\TFG ALZHEIMER\GUIA_DESARROLLO_CUIDALINK.docx"
doc.save(output_path)
print(f"✓ Guía de desarrollo creada: {output_path}")
print(f"✓ Tamaño: Comprensiva (~30+ páginas)")
print(f"✓ Secciones: Arquitectura, Frontend, Backend, BD, Flujos, Git, Testing, Deployment, Troubleshooting")
