#!/usr/bin/env python3
"""
Genera el documento Word con las respuestas a la rúbrica de BBDD y APIs de CuidaLink.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Estilos globales ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Función helper para tablas con bordes ──
def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single', qn('w:sz'): '4',
            qn('w:space'): '0', qn('w:color'): '999999'
        })
        borders.append(el)
    tblPr.append(borders)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml import OxmlElement
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '2E7D32')
        shading.set(qn('w:val'), 'clear')
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        cell._tc.get_or_add_tcPr().append(shading)
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def heading1(text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

def heading2(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x38, 0x8E, 0x3C)

def heading3(text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x43, 0xA0, 0x47)

def para(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
    p.add_run(text)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
    p.add_run(text)
    return p

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

# ═══════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CuidaLink')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Rúbrica: Base de Datos y APIs')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Trabajo Fin de Grado — App de asistencia a cuidadores de pacientes con Alzheimer')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('Marzo 2026')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. JUSTIFICACIÓN DEL MOTOR DE BASE DE DATOS
# ═══════════════════════════════════════════════════════════════
heading1('1. Justificación del motor de base de datos')

para('CuidaLink utiliza una estrategia multi-motor con 3 SGBD según el entorno de ejecución, lo que permite optimizar cada fase del desarrollo:')

add_table(
    ['Entorno', 'Motor', 'Justificación'],
    [
        ['Desarrollo', 'H2 (in-memory)',
         'Base de datos embebida que no requiere instalación. Con ddl-auto: create-drop el esquema se recrea en cada arranque, ideal para iterar rápido sin mantener migraciones. Se carga data-h2.sql automáticamente con datos demo.'],
        ['Default / Local', 'MySQL 8.0+',
         'Motor relacional robusto y ampliamente adoptado. Soporta JSON nativo (usado en datos_extra de eventos y ruta_geojson de paseos), ENUM para tipos de eventos, LONGTEXT para fotos Base64 y DECIMAL(10,8) para coordenadas GPS de alta precisión.'],
        ['Producción', 'PostgreSQL 15+ (Neon Serverless)',
         'Elegido por su despliegue serverless en la nube (Neon), escalado automático y coste cero en plan free. Ofrece CHECK constraints más flexibles que ENUM de MySQL, BIGSERIAL para auto-increment, y mejor concurrencia MVCC para múltiples cuidadoras accediendo simultáneamente.'],
    ]
)

doc.add_paragraph()
heading2('¿Por qué relacional y no NoSQL?')

para('Los datos de CuidaLink son altamente relacionales: un paciente tiene múltiples eventos, cada evento pertenece a una cuidadora, los ratings vinculan familiares con cuidadoras. Las foreign keys con ON DELETE CASCADE/SET NULL garantizan integridad referencial que un motor documental (MongoDB) no ofrece de forma nativa. Además, las consultas temporales (eventos entre fechas, citas por rango) se benefician de los índices B-tree (idx_abuelo_fecha, idx_eventos_tipo).')

heading2('Limitaciones reconocidas')

para('H2 no soporta ENUM ni JSON nativo como MySQL, por lo que se usa VARCHAR con CHECK en PostgreSQL/H2 y TEXT para datos JSON. El almacenamiento de fotos en Base64 dentro de TEXT/LONGTEXT consume más espacio que un almacenamiento en fichero/S3, pero simplifica la arquitectura al no necesitar un servicio de almacenamiento externo para un TFG.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. DISEÑO Y ESTRUCTURA DE LA BASE DE DATOS
# ═══════════════════════════════════════════════════════════════
heading1('2. Diseño y estructura de la base de datos')

para('El esquema consta de 9 tablas con relaciones normalizadas en 3FN (Tercera Forma Normal):')

heading2('Modelo entidad-relación')

code_block(
    '┌──────────┐    1:N    ┌──────────┐    N:1    ┌──────────┐\n'
    '│ Cuidadora├──────────>│  Evento  │<──────────┤  Abuelo  │\n'
    '│          ├──────────>│          │           │          │\n'
    '│  PK: id  │    1:N    │ FK: cuid │           │ FK: fam  │\n'
    '│  tel UNQ │           │ FK: abul │           │  lat/lng │\n'
    '│  pin BCR │           │ tipo ENUM│           └────┬─────┘\n'
    '│  rating  │           │ gps,foto │              1 │\n'
    '└────┬─────┘           └──────────┘                │\n'
    '     │ 1:N                                         │ N\n'
    '┌────┴─────┐                                 ┌─────┴────┐\n'
    '│  Rating  │    N:1                          │ Familiar │\n'
    '│ estrellas├─────────── Abuelo               │ email UNQ│\n'
    '│ FK: cuid │                                 │ pwd BCR  │\n'
    '│ FK: abul │                                 └──────────┘\n'
    '└──────────┘\n'
    '\n'
    '┌──────────┐  ┌──────────┐  ┌──────────┐  ┌───────────┐\n'
    '│  Paseo   │  │   Cita   │  │   Nota   │  │MensajeChat│\n'
    '│ FK: cuid │  │ FK: abul │  │ abueloId │  │ abueloId  │\n'
    '│ FK: abul │  │ fecha+hora│ │ autor    │  │ remitente │\n'
    '│ GeoJSON  │  │ tipo     │  │ prioridad│  │ text,audio│\n'
    '│ distancia│  │ creadoPor│  │ leida    │  │ image     │\n'
    '└──────────┘  └──────────┘  └──────────┘  └───────────┘'
)

heading2('Detalle de las 9 tablas')

heading3('Tabla: abuelos (Pacientes)')
add_table(
    ['Campo', 'Tipo', 'Restricciones', 'Descripción'],
    [
        ['id', 'BIGINT / BIGSERIAL', 'PK, AUTO_INCREMENT', 'Identificador único'],
        ['nombre', 'VARCHAR(100)', 'NOT NULL', 'Nombre del paciente'],
        ['direccion', 'VARCHAR(255)', '', 'Dirección del domicilio'],
        ['lat / lng', 'DECIMAL(10,8) / (11,8)', '', 'Coordenadas GPS del domicilio (geofencing)'],
        ['familiar_id', 'BIGINT', 'FK → familiares', 'Familiar supervisor'],
        ['telefono_emergencia', 'VARCHAR(15)', '', 'Teléfono de emergencia'],
        ['foto_perfil', 'TEXT / LONGTEXT', '', 'Foto en Base64'],
        ['notas_medicas', 'TEXT', '', 'Información médica relevante'],
        ['created_at', 'DATETIME / TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP', 'Fecha de creación'],
    ]
)

doc.add_paragraph()
heading3('Tabla: cuidadoras (Profesionales)')
add_table(
    ['Campo', 'Tipo', 'Restricciones', 'Descripción'],
    [
        ['id', 'BIGINT / BIGSERIAL', 'PK, AUTO_INCREMENT', 'Identificador único'],
        ['nombre', 'VARCHAR(100)', 'NOT NULL', 'Nombre de la cuidadora'],
        ['telefono', 'VARCHAR(15)', 'UNIQUE', 'Identificador de login'],
        ['pin', 'VARCHAR(60)', '', 'PIN hasheado con BCrypt'],
        ['foto_perfil', 'TEXT / LONGTEXT', '', 'Foto en Base64'],
        ['rating_promedio', 'DECIMAL(2,1)', 'DEFAULT 0', 'Media de valoraciones (1-5)'],
        ['total_ratings', 'INT', 'DEFAULT 0', 'Número total de valoraciones'],
        ['activo', 'BOOLEAN', 'DEFAULT TRUE', 'Cuenta activa/desactivada'],
        ['created_at', 'DATETIME / TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP', 'Fecha de alta'],
    ]
)

doc.add_paragraph()
heading3('Tabla: familiares')
add_table(
    ['Campo', 'Tipo', 'Restricciones', 'Descripción'],
    [
        ['id', 'BIGINT / BIGSERIAL', 'PK, AUTO_INCREMENT', 'Identificador único'],
        ['nombre', 'VARCHAR(100)', 'NOT NULL', 'Nombre del familiar'],
        ['email', 'VARCHAR(100)', 'UNIQUE', 'Identificador de login'],
        ['telefono', 'VARCHAR(15)', '', 'Teléfono de contacto'],
        ['password', 'VARCHAR(60)', 'NOT NULL', 'Contraseña hasheada con BCrypt'],
        ['parentesco', 'VARCHAR(50)', '', 'Relación con el paciente (Hijo/a, Nieto/a...)'],
        ['activo', 'BOOLEAN', 'DEFAULT TRUE', 'Cuenta activa/desactivada'],
        ['created_at', 'DATETIME / TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP', 'Fecha de registro'],
    ]
)

doc.add_paragraph()
heading3('Tabla: eventos (Central del sistema)')
add_table(
    ['Campo', 'Tipo', 'Restricciones', 'Descripción'],
    [
        ['id', 'BIGINT / BIGSERIAL', 'PK, AUTO_INCREMENT', 'Identificador único'],
        ['abuelo_id', 'BIGINT', 'FK → abuelos ON DELETE SET NULL', 'Paciente asociado'],
        ['cuidadora_id', 'BIGINT', 'FK → cuidadoras ON DELETE SET NULL', 'Cuidadora que registra'],
        ['tipo', 'ENUM / VARCHAR(20)', 'CHECK IN (LLEGADA, PASTILLA, COMIDA, PASEO, SIESTA, CAIDA, SALIDA, FUGA)', 'Tipo de actividad'],
        ['timestamp', 'DATETIME / TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP', 'Momento del evento'],
        ['foto_base64', 'TEXT / LONGTEXT', '', 'Foto de verificación'],
        ['gps_lat / gps_lng', 'DECIMAL(10,8) / (11,8)', '', 'Ubicación GPS del evento'],
        ['verificado', 'BOOLEAN', 'DEFAULT FALSE', 'Verificado por geofencing/OCR'],
        ['datos_extra', 'JSON / TEXT', '', 'Datos adicionales (nombre medicamento, texto OCR)'],
        ['descripcion', 'VARCHAR(255)', '', 'Descripción libre'],
    ]
)

doc.add_paragraph()
heading3('Tabla: ratings (Valoraciones)')
add_table(
    ['Campo', 'Tipo', 'Restricciones', 'Descripción'],
    [
        ['id', 'BIGINT / BIGSERIAL', 'PK, AUTO_INCREMENT', 'Identificador único'],
        ['cuidadora_id', 'BIGINT', 'FK → cuidadoras ON DELETE CASCADE', 'Cuidadora valorada'],
        ['abuelo_id', 'BIGINT', 'FK → abuelos ON DELETE SET NULL', 'Paciente del contexto'],
        ['familiar_id', 'BIGINT', '', 'Familiar que valora'],
        ['estrellas', 'TINYINT / SMALLINT', 'CHECK BETWEEN 1 AND 5', 'Puntuación 1-5 estrellas'],
        ['comentario', 'TEXT', '', 'Comentario opcional'],
        ['fecha', 'DATETIME / TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP', 'Fecha de la valoración'],
    ]
)

doc.add_paragraph()
heading3('Tabla: paseos (Tracking GPS)')
add_table(
    ['Campo', 'Tipo', 'Restricciones', 'Descripción'],
    [
        ['id', 'BIGINT / BIGSERIAL', 'PK, AUTO_INCREMENT', 'Identificador único'],
        ['abuelo_id', 'BIGINT', 'FK → abuelos', 'Paciente del paseo'],
        ['cuidadora_id', 'BIGINT', 'FK → cuidadoras', 'Cuidadora que acompaña'],
        ['inicio / fin', 'DATETIME / TIMESTAMP', '', 'Inicio y fin del paseo'],
        ['ruta_geojson', 'JSON / TEXT', '', 'Ruta completa en formato GeoJSON LineString'],
        ['distancia_km', 'DECIMAL(4,2)', '', 'Distancia recorrida en km'],
        ['activo', 'BOOLEAN', 'DEFAULT TRUE', 'Paseo en curso'],
    ]
)

doc.add_paragraph()
heading3('Tabla: citas (Calendario compartido)')
add_table(
    ['Campo', 'Tipo', 'Restricciones', 'Descripción'],
    [
        ['id', 'BIGINT / BIGSERIAL', 'PK, AUTO_INCREMENT', 'Identificador único'],
        ['titulo', 'VARCHAR(255)', 'NOT NULL', 'Título de la cita'],
        ['fecha', 'DATE', 'NOT NULL', 'Fecha de la cita'],
        ['hora', 'TIME', 'NOT NULL', 'Hora de la cita'],
        ['tipo', 'ENUM / VARCHAR(20)', 'CHECK IN (CITA_MEDICA, MEDICAMENTO, ACTIVIDAD, VISITA, OTRO)', 'Categoría'],
        ['notas', 'VARCHAR(500)', '', 'Notas adicionales'],
        ['creado_por', 'VARCHAR(20)', 'NOT NULL', '"cuidadora" o "familiar"'],
        ['abuelo_id', 'BIGINT', 'FK → abuelos ON DELETE CASCADE', 'Paciente asociado'],
    ]
)

doc.add_paragraph()
heading3('Tabla: notas')
add_table(
    ['Campo', 'Tipo', 'Restricciones', 'Descripción'],
    [
        ['id', 'BIGINT / BIGSERIAL', 'PK, AUTO_INCREMENT', 'Identificador único'],
        ['texto', 'VARCHAR(1000)', 'NOT NULL', 'Contenido de la nota'],
        ['autor', 'VARCHAR(20)', 'NOT NULL', '"cuidadora" o "familiar"'],
        ['timestamp', 'TIMESTAMP', 'NOT NULL', 'Fecha/hora de creación'],
        ['leida', 'BOOLEAN', 'DEFAULT FALSE', 'Estado de lectura'],
        ['prioridad', 'VARCHAR(10)', 'DEFAULT "normal"', '"normal" o "urgente"'],
        ['abuelo_id', 'BIGINT', 'DEFAULT 1', 'Paciente asociado'],
    ]
)

doc.add_paragraph()
heading3('Tabla: mensajes_chat')
add_table(
    ['Campo', 'Tipo', 'Restricciones', 'Descripción'],
    [
        ['id', 'BIGINT / BIGSERIAL', 'PK, AUTO_INCREMENT', 'Identificador único'],
        ['abuelo_id', 'BIGINT', 'NOT NULL', 'Sala de chat por paciente'],
        ['remitente', 'VARCHAR(20)', 'NOT NULL', '"cuidadora" o "familiar"'],
        ['text', 'VARCHAR(1000)', 'NOT NULL', 'Texto del mensaje'],
        ['type', 'VARCHAR(10)', 'DEFAULT "text"', '"text", "audio" o "image"'],
        ['audio_base64', 'TEXT', '', 'Nota de voz en Base64'],
        ['audio_duration', 'INT', '', 'Duración del audio en segundos'],
        ['image_base64', 'TEXT', '', 'Imagen en Base64'],
        ['timestamp', 'TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP', 'Fecha/hora del mensaje'],
        ['leido', 'BOOLEAN', 'DEFAULT FALSE', 'Estado de lectura'],
    ]
)

doc.add_paragraph()
heading2('Normalización y coherencia')

bullet('Claves primarias: ', 'BIGINT/BIGSERIAL AUTO_INCREMENT')
bullet(' en todas las tablas')
bullet('Claves foráneas con integridad referencial: ', 'ON DELETE CASCADE ')
bullet('(ratings se borran si se elimina la cuidadora), ', 'ON DELETE SET NULL ')
bullet('(eventos conservan historial si se elimina paciente)')
bullet('Índices compuestos para consultas frecuentes: ', 'idx_abuelo_fecha (abuelo_id, timestamp) ')
bullet('optimiza "eventos de hoy por paciente", ', 'idx_paseos_activo (cuidadora_id, activo) ')
bullet('optimiza la búsqueda de paseo activo')
bullet('Tipos de datos específicos: ', 'DECIMAL(10,8)/(11,8) ')
bullet('para coordenadas GPS con precisión milimétrica, ', 'TINYINT CHECK (BETWEEN 1 AND 5) ')
bullet('para ratings')
bullet('Campos created_at con DEFAULT CURRENT_TIMESTAMP y @PrePersist en JPA para auditoría automática', '')
bullet('Constraints UNIQUE en cuidadoras.telefono y familiares.email para evitar duplicados de acceso', '')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. CONEXIÓN DE LA APLICACIÓN CON LA BBDD
# ═══════════════════════════════════════════════════════════════
heading1('3. Conexión de la aplicación con la BBDD')

para('La conexión se gestiona mediante Spring Data JPA + Hibernate + HikariCP (connection pool incluido en Spring Boot por defecto).')

heading2('Configuración por perfiles')

heading3('Perfil de Desarrollo (H2 in-memory)')
code_block(
    'spring:\n'
    '  datasource:\n'
    '    url: jdbc:h2:mem:cuidalink;DB_CLOSE_DELAY=-1\n'
    '    driver-class-name: org.h2.Driver\n'
    '  jpa:\n'
    '    hibernate:\n'
    '      ddl-auto: create-drop\n'
    '  sql:\n'
    '    init:\n'
    '      data-locations: classpath:data-h2.sql'
)

heading3('Perfil de Producción (PostgreSQL Neon Cloud)')
code_block(
    'spring:\n'
    '  datasource:\n'
    '    url: ${DATABASE_URL}                    # Variable de entorno\n'
    '    driver-class-name: org.postgresql.Driver\n'
    '    hikari:\n'
    '      maximum-pool-size: 5                  # Limitado para Neon free tier\n'
    '      connection-timeout: 20000             # 20s timeout\n'
    '  jpa:\n'
    '    hibernate:\n'
    '      ddl-auto: update                      # Migraciones incrementales'
)

heading2('Buenas prácticas implementadas')

bullet('Connection pooling (HikariCP): ', '')
para('Pool de 5 conexiones en producción para no saturar el tier gratuito de Neon, pool por defecto de 10 en desarrollo. HikariCP es el pool más rápido de JVM, incluido automáticamente por Spring Boot.')

bullet('Carga lazy (@ManyToOne(fetch = LAZY)): ', '')
para('Las relaciones Evento→Abuelo y Evento→Cuidadora no cargan datos innecesarios hasta que se accede a ellos, evitando el problema N+1 de consultas.')

bullet('Paginación: ', '')
para('findByAbueloIdOrderByTimestampDesc(Long, Pageable) limita resultados a páginas de 20 (máximo 100) para evitar cargar miles de eventos en memoria.')

bullet('Queries JPQL optimizadas: ', '')
code_block(
    '@Query("SELECT e FROM Evento e WHERE e.abuelo.id = :abueloId\n'
    '        AND DATE(e.timestamp) = CURRENT_DATE\n'
    '        ORDER BY e.timestamp DESC")\n'
    'List<Evento> findTodayEventsByAbueloId(@Param("abueloId") Long abueloId);'
)

bullet('Gestión de errores: ', '')
para('GlobalExceptionHandler captura excepciones de persistencia y devuelve respuestas HTTP apropiadas (400/500) sin exponer stack traces ni información interna de la base de datos.')

bullet('Interceptor de respuesta (cliente móvil): ', '')
para('En el móvil, si el backend devuelve 401 (token expirado), el interceptor Axios limpia automáticamente AsyncStorage y redirige a login.')

heading2('Ejemplo completo de flujo de conexión')

code_block(
    '1. App móvil → POST /api/checkin {selfie, gps}\n'
    '2. JwtAuthFilter valida token → extrae cuidadoraId\n'
    '3. EventoService.registrarCheckin() ejecuta lógica\n'
    '4. Spring Data JPA → HikariCP obtiene conexión del pool\n'
    '5. Hibernate genera INSERT INTO eventos (...) VALUES (?,...)\n'
    '6. PreparedStatement parametrizado (previene SQL injection)\n'
    '7. PostgreSQL/MySQL ejecuta la query\n'
    '8. HikariCP devuelve la conexión al pool\n'
    '9. WebSocketService notifica al familiar en tiempo real\n'
    '10. Controller devuelve 200 + EventoResponse JSON'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. INTEGRACIÓN Y USO DE APIs
# ═══════════════════════════════════════════════════════════════
heading1('4. Integración y uso de APIs')

para('CuidaLink integra una API REST propia desarrollada con Spring Boot y múltiples APIs nativas del dispositivo móvil.')

heading2('4.1. API REST propia (Spring Boot)')

add_table(
    ['Módulo', 'Endpoints', 'Descripción'],
    [
        ['Autenticación', 'POST /auth/login\nPOST /auth/familiar/login\nPOST /auth/familiar/register', 'Login cuidadora (tel+PIN), login familiar (email+pwd), registro'],
        ['Eventos', 'POST /checkin, /pastilla, /comida, /caida\nGET /eventos/{id}/hoy', 'CRUD de actividades diarias con verificación GPS y fotos'],
        ['Dashboard', 'GET /dashboard/{id}/hoy', 'Agregación: datos paciente + eventos + tareas + % completado'],
        ['Chat', 'POST /chat/send\nGET /chat/messages\nDELETE /chat/{id}', 'Mensajería texto/audio/imagen entre cuidadora y familiar'],
        ['Paseos GPS', 'POST /paseo/start\nPOST /paseo/stop', 'Tracking de rutas en formato GeoJSON'],
        ['Calendario', 'CRUD /citas + filtros fecha/rango', 'Citas compartidas entre cuidadora y familiar'],
        ['Notas', 'CRUD /notas + /pendientes', 'Notas con estado leída/no leída y prioridad'],
        ['Valoraciones', 'POST /rating\nGET /rating/{id}', 'Sistema 1-5 estrellas con cálculo de promedio'],
        ['Admin', 'CRUD bajo /auth/admin/*', 'Gestión de cuentas de cuidadoras y familiares'],
    ]
)

doc.add_paragraph()
heading2('4.2. Integración desde el móvil — Cliente Axios')

code_block(
    '// api.ts — Cliente HTTP con interceptores JWT automáticos\n'
    'const api = axios.create({ baseURL: API_URL, timeout: 3000 });\n'
    '\n'
    '// Interceptor REQUEST: inyecta token en cada petición\n'
    'api.interceptors.request.use(async (config) => {\n'
    '    const token = await AsyncStorage.getItem("token");\n'
    '    if (token) config.headers.Authorization = `Bearer ${token}`;\n'
    '    return config;\n'
    '});\n'
    '\n'
    '// Interceptor RESPONSE: limpia sesión en 401\n'
    'api.interceptors.response.use(response => response, async (error) => {\n'
    '    if (error.response?.status === 401) {\n'
    '        await AsyncStorage.removeItem("token");\n'
    '    }\n'
    '    return Promise.reject(error);\n'
    '});'
)

doc.add_paragraph()
heading2('4.3. APIs nativas del dispositivo')

add_table(
    ['API', 'Paquete Expo', 'Uso en CuidaLink'],
    [
        ['GPS / Geolocalización', 'expo-location', 'Check-in con geofencing (200m), tracking de paseos, alertas de fuga'],
        ['Cámara', 'expo-camera', 'Selfies de llegada, fotos de medicamentos (OCR), fotos de comida'],
        ['Acelerómetro', 'expo-sensors', 'Detección de caídas (umbral de aceleración)'],
        ['Audio', 'expo-audio', 'Notas de voz en el chat'],
        ['Notificaciones', 'expo-notifications', 'Alertas de eventos, recordatorios de medicación'],
        ['WebSocket/STOMP', 'Spring Backend', 'Notificaciones en tiempo real: eventos, emergencias, mensajes chat'],
    ]
)

doc.add_paragraph()
heading2('4.4. Diagrama de integración')

code_block(
    '┌─────────────────┐     HTTPS/JWT      ┌─────────────────┐     JDBC/HikariCP\n'
    '│  React Native   │ <────────────────> │  Spring Boot    │ <────────────────>  PostgreSQL\n'
    '│  (Expo SDK 54)  │     REST API       │  (Java 17)      │                    (Neon Cloud)\n'
    '│                 │                    │                 │\n'
    '│  Axios Client   │     WebSocket      │  STOMP Broker   │\n'
    '│  + JWT Intercep │ <────────────────> │  /topic/familia │\n'
    '└─────────────────┘                    └─────────────────┘\n'
    '     │\n'
    '     ├── expo-location (GPS)\n'
    '     ├── expo-camera (Fotos)\n'
    '     ├── expo-sensors (Acelerómetro)\n'
    '     ├── expo-audio (Voz)\n'
    '     └── expo-notifications (Push)'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. SEGURIDAD EN BASE DE DATOS Y APIs
# ═══════════════════════════════════════════════════════════════
heading1('5. Seguridad en base de datos y APIs')

heading2('5.1. Autenticación (quién eres)')

add_table(
    ['Rol', 'Método', 'Implementación'],
    [
        ['Cuidadora', 'Teléfono + PIN 4 dígitos', 'PIN hasheado con BCrypt ($2a$10$...), verificado con BCryptPasswordEncoder.matches()'],
        ['Familiar', 'Email + Contraseña', 'Contraseña hasheada con BCrypt, verificación idéntica'],
    ]
)

doc.add_paragraph()
para('Las contraseñas nunca se almacenan en texto plano. Los campos pin (60 caracteres) y password (60 caracteres) almacenan el hash BCrypt con salt aleatorio, lo que hace que incluso contraseñas iguales generen hashes diferentes.')

heading2('5.2. Autorización (qué puedes hacer)')

bullet('JWT con roles: Cada token contiene role: "CUIDADORA" o "FAMILIAR", extraído por JwtTokenProvider.getRoleFromToken()', '')
bullet('Rutas protegidas:', '')
code_block(
    '// SecurityConfig.java\n'
    '.requestMatchers("/api/auth/**").permitAll()    // Login/registro: público\n'
    '.anyRequest().authenticated()                    // Todo lo demás: requiere JWT válido'
)
bullet('Sesiones stateless: SessionCreationPolicy.STATELESS — no se usan cookies de sesión, cada request lleva su JWT', '')
bullet('CSRF deshabilitado: Apropiado para APIs REST stateless (no hay formularios HTML)', '')

heading2('5.3. Tokens JWT')

add_table(
    ['Propiedad', 'Valor'],
    [
        ['Algoritmo', 'HMAC-SHA256 (secreto de 256+ bits)'],
        ['Librería', 'jjwt 0.12.3 (io.jsonwebtoken)'],
        ['Expiración', '24 horas (86.400.000 ms), configurable'],
        ['Claims', 'sub (userId), identifier (teléfono/email), role (CUIDADORA/FAMILIAR)'],
        ['Filtro', 'JwtAuthenticationFilter (extends OncePerRequestFilter)'],
        ['Interceptor cliente', 'Axios: Bearer token automático en cada request; limpia AsyncStorage en 401'],
    ]
)

doc.add_paragraph()
heading2('5.4. Protección de la base de datos')

bullet('Contraseñas hasheadas: BCrypt con salt aleatorio (cost factor 10)', '')
bullet('SQL Injection prevenida: Spring Data JPA usa PreparedStatements parametrizado en todas las queries — tanto las derivadas del nombre de método (findByEmail) como las @Query JPQL', '')
bullet('Gestión de errores: GlobalExceptionHandler nunca expone stack traces ni información interna de la BBDD al cliente', '')
bullet('Variables de entorno en producción: DATABASE_URL, DATABASE_USER, DATABASE_PASSWORD, JWT_SECRET se inyectan como variables de entorno, nunca hardcodeadas en el código', '')
bullet('Connection pooling limitado: HikariCP con maximum-pool-size: 5 en producción para prevenir DoS por agotamiento de conexiones', '')
bullet('CORS restringido: Solo orígenes permitidos (localhost:3000, localhost:5173, etc.) pueden acceder a la API', '')

heading2('5.5. Diagrama de flujo de seguridad')

code_block(
    '┌──────────┐    email/pwd     ┌──────────────┐   BCrypt.matches()  ┌──────────┐\n'
    '│  Móvil   │ ───────────────> │ POST /login   │ ──────────────────>│   BBDD   │\n'
    '│          │                  │               │                    │ hash pwd │\n'
    '│          │ <─────────────── │ JWT generado  │ <──────────────────│ ✓ válido │\n'
    '│          │    { token }     └───────────────┘                    └──────────┘\n'
    '│          │\n'
    '│          │    GET /eventos\n'
    '│          │    Authorization: Bearer <JWT>\n'
    '│          │ ───────────────> ┌──────────────────┐\n'
    '│          │                  │ JwtAuthFilter     │\n'
    '│          │                  │ validateToken()   │\n'
    '│          │                  │ setAuthentication │\n'
    '│          │ <─────────────── │ → Controller      │\n'
    '│          │   200 + datos    └──────────────────┘\n'
    '└──────────┘'
)

# ── Guardar ──
output_path = os.path.join(os.path.dirname(__file__), 'RUBRICA_BBDD_APIS_CUIDALINK.docx')
doc.save(output_path)
print(f'Documento generado: {output_path}')
